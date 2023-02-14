from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import datetime

now_time = datetime.datetime.now()
TodayDate = now_time.strftime('%Y/%m/%d')
TodayFile = now_time.strftime('%Y_%m%d' + '.tmp')
filename = now_time.strftime('%Y_%m%d' + '.docx')

data_delimiter = '(\D)'
section_delimiter = '(\S)'
indent_block_delimiter = '(\I)'

f = open(TodayFile,'r')
textdata = f.read()
f.close()

pretext = textdata.split(data_delimiter)

RCOSID = pretext[0]
textdata = pretext[1]

doc = Document()
header_section = doc.sections[0].header
ToDate = header_section.paragraphs[0]
ToDate.text = str(TodayDate) + RCOSID
ToDate.alignment = WD_TAB_ALIGNMENT.RIGHT

doc.add_paragraph('')
sections = textdata.split(section_delimiter)
for item in sections:
  subsections = item.split(indent_block_delimiter)
  for subitem in subsections:
    doc.add_paragraph(subitem)
  #doc.add_paragraph(item)
doc.paragraphs[1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(filename)

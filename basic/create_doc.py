from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import datetime

now_time = datetime.datetime.now()
TodayDate = now_time.strftime('%Y/%m/%d')
TodayFile = now_time.strftime('%Y_%m%d' + '.tmp')
filename = now_time.strftime('%Y_%m%d' + '.docx')

section_delimiter = '(\S)'

f = open(TodayFile,'r')
textdata = f.read()
f.close()

RCOSID = '[R-2022-XX]'

doc = Document()
header_section = doc.sections[0].header
ToDate = header_section.paragraphs[0]
ToDate.text = str(TodayDate) + '\n' + RCOSID
ToDate.alignment = WD_TAB_ALIGNMENT.RIGHT

doc.add_paragraph('')
sections = textdata.split(section_delimiter)
for item in sections:
  doc.add_paragraph(item)
#doc.add_paragraph(sections[0])
doc.paragraphs[1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(filename)
